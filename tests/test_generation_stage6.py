"""Testes da Etapa 6 — generate(artifact_type, deal_state)."""

from __future__ import annotations

import hashlib
import os
from pathlib import Path

import pytest

from app.artifact_types import REQUIRED_SECTIONS, resolve_artifact_type
from app.generation.citations import collect_citations
from app.generation.content import build_document_sections
from app.generation.generate import UnsupportedArtifactTypeError, generate
from app.repositories.workspace_artifact import compute_file_hash


def _deal_state(**overrides) -> dict:
    base = {
        "job_id": "test-job-stage6",
        "company_name": "Loggi",
        "document_type": "CIM",
        "research_structured": {
            "company_name": "Loggi",
            "sector": "Logística",
            "description": "Plataforma de entregas urbanas.",
            "business_model": "Marketplace B2B",
            "market_position": "Líder",
            "main_competitors": ["Rappi", "iFood"],
        },
        "financial_structured": {
            "revenue_2024": "R$ 2 bi est.",
            "ebitda": "15% est.",
            "valuation": "R$ 5 bi est.",
        },
        "financial_citations": [
            {
                "source_file": "teaser.pdf",
                "page": 3,
                "chunk_id": "c1",
                "quote": "Receita 2024 estimada em R$ 2 bi.",
                "source": "data_room",
            }
        ],
        "research_citations": [],
    }
    base.update(overrides)
    return base


def test_resolve_artifact_type_cim_vs_memo():
    assert resolve_artifact_type("CIM") == "cim_pptx"
    assert resolve_artifact_type("CIM", message="gerar memo de investimento") == "memo_docx"
    assert resolve_artifact_type("VALUATION") == "memo_docx"


def test_shared_content_sections_differ_by_artifact_type():
    state = _deal_state()
    cim = build_document_sections(state, "cim_pptx")
    memo = build_document_sections(state, "memo_docx")
    assert "Resumo Executivo" in cim
    assert "Sumário Executivo" in memo
    assert set(cim.keys()) != set(memo.keys())


def test_citations_shared_across_formats():
    bundle = collect_citations(_deal_state())
    assert bundle.has_citations
    assert bundle.reference_lines[0].startswith("[1]")
    assert "data_room" in bundle.reference_lines[0]


def test_generate_memo_docx_with_citations(tmp_path, monkeypatch):
    out_dir = tmp_path / "outputs"
    out_dir.mkdir()
    monkeypatch.chdir(tmp_path)

    state = _deal_state(document_type="MEMO", artifact_type="memo_docx")
    result = generate("memo_docx", state)

    assert result.format == "docx"
    assert result.artifact_type == "memo_docx"
    assert Path(result.file_path).is_file()
    assert Path(result.file_path).stat().st_size > 500

    from docx import Document

    doc = Document(result.file_path)
    headings = [p.text for p in doc.paragraphs if p.style.name.startswith("Heading")]
    for section in REQUIRED_SECTIONS["memo_docx"]:
        if section == "Referências":
            continue
        assert section in headings
    assert any("teaser.pdf" in p.text for p in doc.paragraphs)


def test_generate_rejects_unsupported_type():
    with pytest.raises(UnsupportedArtifactTypeError):
        generate("model_xlsx", _deal_state())


def test_file_hash_is_bytes_only(tmp_path):
    f = tmp_path / "artifact.bin"
    f.write_bytes(b"conteudo-do-artefato")
    h1 = compute_file_hash(str(f))
    h2 = hashlib.sha256(b"conteudo-do-artefato").hexdigest()
    assert h1 == h2

    # Audits não alteram file_hash
    assert compute_file_hash(str(f)) == h1


def test_qa_sections_parametrized_by_artifact_type(tmp_path, monkeypatch):
    from app.agents.qa_factual import validate_sections

    monkeypatch.chdir(tmp_path)
    gen = generate("memo_docx", {**_deal_state(), "job_id": "qa-memo-test"})
    state = _deal_state()
    state["artifact_type"] = "memo_docx"
    state["file_path"] = gen.file_path
    issues = validate_sections(state)
    assert not any("slides" in i.lower() for i in issues)
