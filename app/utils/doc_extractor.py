"""Extração de texto de documentos em memória — nunca persiste arquivos."""

from __future__ import annotations

import csv
import io
from typing import Tuple

MAX_EXTRACTED_CHARS = 120_000


def _truncate(text: str) -> str:
    text = text.strip()
    if len(text) <= MAX_EXTRACTED_CHARS:
        return text
    return text[:MAX_EXTRACTED_CHARS] + "\n\n[... conteúdo truncado ...]"


def extract_pdf(data: bytes) -> str:
    import fitz

    doc = fitz.open(stream=data, filetype="pdf")
    parts: list[str] = []
    try:
        for page in doc:
            parts.append(page.get_text())
    finally:
        doc.close()
    return _truncate("\n".join(parts))


def extract_excel(data: bytes, filename: str = "") -> str:
    name = (filename or "").lower()
    if name.endswith(".xls") and not name.endswith(".xlsx"):
        return _extract_xls(data)
    return _extract_xlsx(data)


def _extract_xlsx(data: bytes) -> str:
    from openpyxl import load_workbook

    wb = load_workbook(io.BytesIO(data), read_only=True, data_only=True)
    parts: list[str] = []
    try:
        for sheet in wb.worksheets:
            parts.append(f"=== Planilha: {sheet.title} ===")
            for row in sheet.iter_rows(values_only=True):
                cells = [str(c).strip() if c is not None else "" for c in row]
                if any(cells):
                    parts.append(" | ".join(cells))
    finally:
        wb.close()
    return _truncate("\n".join(parts))


def _extract_xls(data: bytes) -> str:
    import xlrd

    book = xlrd.open_workbook(file_contents=data)
    parts: list[str] = []
    for sheet in book.sheets():
        parts.append(f"=== Planilha: {sheet.name} ===")
        for row_idx in range(sheet.nrows):
            cells = [str(sheet.cell_value(row_idx, col_idx)).strip() for col_idx in range(sheet.ncols)]
            if any(cells):
                parts.append(" | ".join(cells))
    return _truncate("\n".join(parts))


def extract_docx(data: bytes) -> str:
    from docx import Document

    doc = Document(io.BytesIO(data))
    parts: list[str] = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            parts.append(text)
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            if any(cells):
                parts.append(" | ".join(cells))
    return _truncate("\n".join(parts))


def extract_csv(data: bytes) -> str:
    for encoding in ("utf-8-sig", "utf-8", "latin-1"):
        try:
            text = data.decode(encoding)
            break
        except UnicodeDecodeError:
            continue
    else:
        text = data.decode("utf-8", errors="replace")

    reader = csv.reader(io.StringIO(text))
    parts: list[str] = []
    for row in reader:
        if any(cell.strip() for cell in row):
            parts.append(" | ".join(cell.strip() for cell in row))
    return _truncate("\n".join(parts))


def _detect_doc_type(filename: str, mime_type: str = "") -> str | None:
    name = (filename or "").lower()
    mime = (mime_type or "").lower()

    if name.endswith(".pdf") or mime == "application/pdf":
        return "pdf"
    if name.endswith(".xlsx") or "spreadsheetml" in mime:
        return "excel"
    if name.endswith(".xls"):
        return "excel"
    if name.endswith(".docx") or "wordprocessingml" in mime:
        return "docx"
    if name.endswith(".csv") or mime in ("text/csv", "application/csv"):
        return "csv"
    return None


def extract_document(data: bytes, filename: str, mime_type: str = "") -> Tuple[str, str]:
    """Extrai texto de bytes. Retorna (texto, doc_type)."""
    doc_type = _detect_doc_type(filename, mime_type)
    if not doc_type:
        raise ValueError(f"Tipo de documento não suportado: {filename or mime_type}")

    if doc_type == "pdf":
        text = extract_pdf(data)
    elif doc_type == "excel":
        text = extract_excel(data, filename)
    elif doc_type == "docx":
        text = extract_docx(data)
    else:
        text = extract_csv(data)

    if not text.strip():
        raise ValueError("Não foi possível extrair texto do documento.")

    return text, doc_type
