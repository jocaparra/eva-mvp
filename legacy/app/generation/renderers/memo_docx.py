"""Renderer investment memo → DOCX."""

from __future__ import annotations

import os
from datetime import date

from app.artifact_types import ARTIFACT_EXTENSION, ARTIFACT_MIME
from app.generation.citations import CitationBundle
from app.generation.models import GenerationResult


class MemoDocxRenderer:
    artifact_type = "memo_docx"

    def render(
        self,
        deal_state: dict,
        sections: dict[str, str],
        citations: CitationBundle,
        output_path: str,
    ) -> GenerationResult:
        from docx import Document

        company = deal_state.get("company_name", "empresa")
        doc_type = deal_state.get("document_type", "MEMO")

        os.makedirs(os.path.dirname(output_path) or ".", exist_ok=True)
        document = Document()
        document.add_heading(f"Investment Memo — {company}", level=0)

        for title, body in sections.items():
            document.add_heading(title, level=1)
            for paragraph in str(body).split("\n"):
                if paragraph.strip():
                    document.add_paragraph(paragraph.strip())

        if citations.has_citations:
            document.add_heading(citations.references_heading(), level=1)
            for line in citations.reference_lines:
                document.add_paragraph(line, style="List Bullet")

        document.save(output_path)
        filename = f"{company}_{doc_type}_{date.today()}{ARTIFACT_EXTENSION[self.artifact_type]}"

        return GenerationResult(
            artifact_type=self.artifact_type,
            file_path=output_path,
            file_filename=filename,
            mime_type=ARTIFACT_MIME[self.artifact_type],
            format="docx",
        )
